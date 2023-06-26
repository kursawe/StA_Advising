import os
import openpyxl
import pandas as pd
import termcolor
from .student import *
import docx
import pathlib


module_catalogue_location = os.path.join(os.path.dirname(__file__),'..','module_catalogue','Module_catalogue.xlsx') 
module_catalogue = pd.read_excel(module_catalogue_location)

def process_form_file(filename):
    """preforms all advising checks on the 
    submitted form.
    
    Parameters:
    -----------
    
    filename : string
        path to the file that is being investigated,
        i.e. a filled-in module choice form
    """ 
    student_or_warning = parse_excel_form(filename)
    _,filename_for_output = os.path.split(filename)
    if isinstance(student_or_warning, str):
        if student_or_warning == 'No student ID':
            warning_message = 'Could not process ' + filename_for_output + '. The file does not contain a valid student ID.'
        elif student_or_warning.startswith('contains invalid student ID'):
            warning_message = 'Could not process ' + filename_for_output + '. The file ' + student_or_warning
        elif student_or_warning.startswith('Do not recognise student programme for parsing:'):
            warning_message = 'Could not process ' + filename_for_output + '. ' + student_or_warning
        colour_code_print_statement(warning_message)
        summary_data =[000000000, 
                    'Unknown',
                    'Unknown',
                    0,
                    warning_message, ' ', ' ', ' ', ' ']

        summary_data_frame = generate_summary_data_frame_from_entries(summary_data)

        return summary_data_frame
    
    student = student_or_warning
    print('Processing file ')
    print(filename)
    print(' ')
    print('Student ID: ' + str(student.student_id))
    print('Name: ' + student.full_name)
    print('Programme: ' + student.programme_name)
    print(' ')
    
    missed_programme_requirements, programme_adviser_recommendations = find_missing_programme_requirements(student)
    
    print('The student is missing the following programme requirements:')
    colour_code_print_statement(missed_programme_requirements)

    missed_prerequisites, prerequisite_adviser_recommendations = find_missing_prerequisites(student)
    
    print('The student is missing the following prerequisites:')
    colour_code_print_statement(missed_prerequisites)

    not_running_modules, scheduling_adviser_recommendations = find_not_running_modules(student)

    print('The student selected the following modules when they are not running:')
    colour_code_print_statement(not_running_modules)

    timetable_clashes, timetable_adviser_recommendations = find_timetable_clashes(student)
    print('I found the following timetable clashes:')
    colour_code_print_statement(timetable_clashes)

    adviser_recommendations = merge_list_to_long_string([programme_adviser_recommendations, prerequisite_adviser_recommendations, 
                                                         scheduling_adviser_recommendations, timetable_adviser_recommendations])

    print('I have the following comments to the adviser:')
    colour_code_print_statement(adviser_recommendations, is_advice = True)

    summary_data = [student.student_id, 
                    student.full_name,
                    student.programme_name,
                    student.current_honours_year,
                    missed_programme_requirements, missed_prerequisites, not_running_modules, timetable_clashes, adviser_recommendations]

    summary_data_frame = generate_summary_data_frame_from_entries(summary_data)
    
    return summary_data_frame

def parse_excel_form(filename):
    """returns an instance of a 'student' class
    that has all the excel data as named attributes

    Parameters:
    -----------
    
    filename : string
        path to the file that is being investigated,
        i.e. a filled-in module choice form
        
    Returns:
    --------

    student : instance of Student class or string
        an object with student attributes. If it's a string then it's a warning
        message about what went wrong
    """
    # open the form file and read in the student ID
    this_workbook = openpyxl.load_workbook(filename=filename)
    sheet = this_workbook.active
    student_id = sheet["D5"].value
    if not isinstance(student_id, int):
        return 'No student ID'
        
    
    # Now that we have the student ID we can look up the student in the database:
    current_directory = os.path.dirname(os.path.abspath(__file__))
    potential_data_files = os.listdir(os.path.join(current_directory,'..', 'student_data'))
    data_files = []
    for candidate_filename in potential_data_files:
       this_filename, file_extension = os.path.splitext(candidate_filename)
       if file_extension == '.csv':
           data_files.append(candidate_filename)
    
    # turn them all into data base files
    data_bases = []
    for data_file_name in data_files:
        data_path = os.path.join(current_directory,'..', 'student_data', data_file_name)
        this_data_frame = pd.read_csv(data_path)
        data_bases.append(this_data_frame)
    
    # get a table with only the entries for this student
    student_not_yet_found = True

    data_base_index = 0
    while student_not_yet_found and data_base_index < len(data_bases):
        this_data_base = data_bases[data_base_index]
        if student_id in this_data_base['Student ID'].to_numpy():
            student_not_yet_found = False
            student_data_base = this_data_base[this_data_base['Student ID'] == student_id]
        data_base_index += 1
    
    if student_not_yet_found:
        return 'contains invalid student ID ' + str(student_id)

    # infer the year of study from the earliest module taken
    data_of_module_years = student_data_base['Year'].str.slice(0,4).astype('int')
    earliest_year = data_of_module_years.min()
    
    #students who took their first module in 2021 will now be in year 3, i.e. 2023-2021+1
    year_of_study = 2023 - earliest_year + 1
    year_of_study = year_of_study
    
    # identify all modules that the student has passed
    data_base_of_passed_modules = student_data_base[student_data_base['Assessment result']=='P']
    passed_modules = data_base_of_passed_modules['Module code'].to_list()
    passed_modules = passed_modules
    
    #identify the programme of the student
    programme_entries = student_data_base['Programme name'].unique()    
    assert(len(programme_entries) == 1)
    programme_name = programme_entries[0]
    
    given_name_entries = student_data_base['Given names'].unique()
    assert(len(given_name_entries) == 1)
    given_name = given_name_entries[0]
    
    family_name_entries = student_data_base['Family name'].unique()
    assert(len(family_name_entries) == 1)
    family_name = family_name_entries[0]
    
    full_name = given_name + ' ' + family_name

    email_entries = student_data_base['Email'].unique()
    assert(len(email_entries) == 1)
    email = email_entries[0]

    # Figure out what year they are in and how many they have left
    if 'Bachelor of Science' in programme_name:
        no_of_programme_years = 4
        expected_honours_years = 2
    elif 'Master in Mathematics' in programme_name:
        no_of_programme_years = 5
        expected_honours_years = 3
    elif 'Master of Arts (Honours)' in programme_name: 
        no_of_programme_years = 4
        expected_honours_years = 2
    elif 'Master in Chemistry (Honours) Chemistry with Mathematics' == programme_name:
        no_of_programme_years = 5
        expected_honours_years = 3
    elif 'Master in Physics (Honours) Mathematics and Theoretical Physics' == programme_name:
        no_of_programme_years = 5
        expected_honours_years = 3
    else:
        warning_message = 'Do not recognise student programme for parsing: ' + programme_name
        return warning_message
    
    if 'EXA120' in student_data_base['Module code'].values:
        no_of_programme_years -=1
        
    no_subhonours_years = no_of_programme_years - expected_honours_years
    current_honours_year = year_of_study - no_subhonours_years
    
    # make a separate data base of passed honours modules
    passed_honours_modules = list()
    for previous_honours_year in range(1,current_honours_year):
        year_difference = current_honours_year - previous_honours_year  
        year_number = 23-year_difference
        calendar_year_string = '20' + str(year_number) + '/20' + str(year_number + 1)
        data_base_of_passed_modules_this_year = data_base_of_passed_modules[data_base_of_passed_modules['Year']==calendar_year_string]
        passed_modules_this_hear = data_base_of_passed_modules_this_year['Module code'].to_list()
        passed_honours_modules += passed_modules_this_hear
    
    # read in modules for all honours years that have not happened yet
    module_table = []
    
    for remaining_honours_year in range(current_honours_year,expected_honours_years + 1):
        year_key = 'Year ' + str(remaining_honours_year)
        calendar_year = 23 + remaining_honours_year - current_honours_year
        calendar_year_string = '20' + str(calendar_year) + '/20' + str(calendar_year + 1)
        for semester_number in [1,2]:
            semester_modules = get_modules_under_header(sheet, year_key + ' of Honours: Semester ' + str(semester_number)) 
            for module in semester_modules:
                module_table.append([year_key, calendar_year_string, 'S' + str(semester_number), module,])

    # make a table for modules that have been passed
    passed_module_table = []
    
    for passed_module in passed_modules:
        module_data_row = student_data_base[student_data_base['Module code'] == passed_module]
        academic_year = module_data_row['Year'].values[0]
        semester = module_data_row['Semester'].values[0]
        this_year = int(academic_year[2:4])
        this_honours_year =  this_year - 23 + current_honours_year
        this_honours_year_string = 'Year ' + str(this_honours_year)
        passed_module_table.append([this_honours_year_string, 
                                    academic_year, semester, passed_module,])

    # Turn this all into a nice pandas data frame
    honours_module_choices = pd.DataFrame(module_table, columns = ['Honours year', 'Academic year', 'Semester', 'Module code'])
    
    passed_module_table = pd.DataFrame(passed_module_table, 
                                       columns = ['Honours year', 'Academic year', 'Semester', 'Module code'])

    this_student = Student(student_id, 
                           full_name,
                           email,
                           programme_name, 
                           year_of_study,
                           expected_honours_years,
                           current_honours_year,
                           passed_modules,
                           passed_module_table,
                           passed_honours_modules,
                           honours_module_choices)
    
    # return the student
    return this_student

def generate_summary_data_frame_from_entries(data_list):
    """Generates the summary data frame from the entries for one student.
    
    Parameters:
    -----------

    data_list : list
        list entries vary in type, see the expected order of entries in the code below.
        
    Returns :
    ---------
    
    summary_data_frame : pandas data frame
        the pandas data frame containing the list data and the correct headers.
    """
    summary_data_frame = pd.DataFrame([data_list], columns = ['Student ID', 
                                                                 'Name',
                                                                 'Programme',
                                                                 'Hon. year',
                                                                 'Unmet programme requirements', 'Missing prerequisites', 'Modules not running', 'Timetable clashes', 'Adviser recommendations'])
    
    return summary_data_frame
   
def colour_code_print_statement(print_statement, is_advice = False):
    """Prints the given string in red if the string is not equal to 'None'.
    Otherwise it will print in green.
    
    Parameters:
    -----------
    
    print_statement : string
        the statement to be printed
        
    is_advice : bool
        if False, strings that are not 'None' will be printed in light magenta instead of red
    """
    if print_statement != 'None':
        if is_advice:
            print(termcolor.colored(print_statement,'blue'))
        else:
            print(termcolor.colored(print_statement,'red'))
    else:
        print(termcolor.colored('None','green'))
        
    #add an empty line by default
    print(' ')

def merge_list_to_long_string(a_list):
    """takes a list of strings and returns a string that separates the entries with a comma and a space.
        Returns the string 'None' if the list is empty. If the first entry of the list is 'None' then that entry will be ignored.
    
    Parameters:
    -----------
    a_list : list
        the list we want to parse, for example a list of unmet programme requirements
    
    Returns:
    --------
    a_string : string
        contains all entries in a_list separated by a comma and a space
        is 'None' if a_list is empty
    """
    a_string = ''
    first_item_found = False
    for item in a_list:
        if item != 'None':
            if first_item_found:
                a_string += ', ' + item
            else:
                a_string = item
                first_item_found = True
    
    if a_string == '':
        a_string = 'None'
    
    return a_string
 
def get_modules_under_header(sheet, header):
    """get all the modules in the student module choice form under a given heading
    
    Parameters:
    -----------
    
    sheet : openpyxl sheet object
        generated from the workbook excel file
        
    header : string
        the header that we are investigating

    Returns:
    --------

    modules : list of strings
        module codes under the given header
    """
    modules = []
    for row in sheet.iter_rows():
        if isinstance(row[1].value, str):
            if header in row[1].value:
                row_number = row[1].row

    # the row number where modules are entered is 2 down
    row_number += 2
    next_cell_name_with_module = 'B' + str(row_number)
    module_code_is_not_empty = sheet[next_cell_name_with_module].value is not None

    while module_code_is_not_empty:
        module_code_entry = sheet[next_cell_name_with_module].value
        if isinstance(module_code_entry, int):
            module_code = str(module_code_entry)
            module_code = 'MT' + module_code
        else:
            module_code = module_code_entry
        module_code = module_code.strip()
        modules.append(module_code)
        row_number+=1 
        next_cell_name_with_module = 'B' + str(row_number)
        module_code_is_not_empty = sheet[next_cell_name_with_module].value is not None

    return modules

def colour_code_passes(column):
    """Helper function for saving the summary data frame, will colour code the column entries in a 
    pandas array in green or red depending on whether they contain the entry 'None' or not
    
    Parameters:
    -----------

    column: pandas data array column
       the column we are colour coding
    
    Returns:
    --------

    color_list : list of strings
       a list of color setting for each of the cells in the array
    """
    passed_colour = 'background-color: palegreen;'
    failed_colour = 'background-color: lightcoral;'
    return [passed_colour if value=='None' else failed_colour for value in column]

def colour_recommendations(column):
    """Helper function for saving the summary data frame, will colour code the column entries in a 
    pandas array in no color or orange depending on whether they contain the entry 'None' or not
    
    Parameters:
    -----------

    column: pandas data array column
       the column we are colour coding
    
    Returns:
    --------

    color_list : list of strings
       a list of color setting for each of the cells in the array
    """
    recommendation_colour = 'background-color: orange;'
    default_colour = ''
    return [default_colour if value=='None' else recommendation_colour for value in column]

def save_summary_data_frame(data_frame, filename):
    """Save the summary data frame two an excel file, inlcuding a colour code formatting 
    for the column.
    
    Parameters :
    ------------

    data_frame : Pandas data frame,
        typically generated from the 'check_form_file' function
    
    filename : string
        where we want to save
    """

    if not filename.endswith('.xlsx'):
        if filename.endswith('.docx'):
            filename = filename[:-5] + '.xlsx'
        else:
            filename += '.xlsx'
        
    styled_data_frame = (data_frame.style.apply(colour_code_passes, subset = ['Unmet programme requirements', 'Missing prerequisites', 'Modules not running',
                                                                                       'Timetable clashes'], axis = 0).
                          apply(colour_recommendations, subset = ['Adviser recommendations'], axis = 0))

    writer = pd.ExcelWriter(filename) 
    # Manually adjust the width of each column
    styled_data_frame.to_excel(writer)
    worksheet = writer.sheets['Sheet1']
    font = openpyxl.styles.Font(size=14)
    worksheet.set_column(0,0,width=5)
    worksheet.set_column(1,1,width=12)
    worksheet.set_column(2,2,width=22)
    worksheet.set_column(3,3,width=35)
    worksheet.set_column(4,4,width=8)
    worksheet.set_column(5,5,width=40)
    worksheet.set_column(6,6,width=40)
    worksheet.set_column(7,7,width=40)
    worksheet.set_column(8,8,width=40)
    worksheet.set_column(9,9,width=40)
    
    # The only way I found to change the font size is to save the excel file, reload it, and then edit the fontsize, and save it again
    writer.save()
    
    # reload
    reloaded_workbook = openpyxl.load_workbook(filename)

    # Access the active sheet
    worksheet = reloaded_workbook.active

    # Set the font size for all cells in the worksheet
    font_size = 14
    font = openpyxl.styles.Font(size=font_size)
    for row in worksheet.iter_rows():
        for cell in row:
            cell.font = font

    # Make the first row bold
    bold_font = openpyxl.styles.Font(size=font_size, bold=True)
    for cell in worksheet[1]:
        cell.font = bold_font
    # Save the modified workbook
    reloaded_workbook.save(filename)
    
    # Now let's also save into word
    filename_base = filename[:-5]
    word_file_name = filename_base + '.docx'
    
    word_document = docx.Document()
    
    for _,row in data_frame.iterrows():
        paragraph = word_document.add_paragraph()
        run = paragraph.add_run('Student ID: ' + str(row['Student ID']))
        run.add_break()
        run = paragraph.add_run('Name: ' + row['Name'])
        run.add_break()
        run = paragraph.add_run('Programme: ' + row['Programme'])
        run.add_break()
        run = paragraph.add_run('The student is missing the following programme requirements:')
        run.add_break()
        run = paragraph.add_run(row['Unmet programme requirements'])
        if row['Unmet programme requirements']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,0,0)
        else:
            run.font.color.rgb = docx.shared.RGBColor(0,128,0)
        run.add_break()
        run = paragraph.add_run('The student is missing the following prerequisites:')
        run.add_break()
        run = paragraph.add_run(row['Missing prerequisites'])
        if row['Missing prerequisites']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,0,0)
        else:
            run.font.color.rgb = docx.shared.RGBColor(0,128,0)
        run.add_break()
        run = paragraph.add_run('The student selected the following modules when they are not running:')
        run.add_break()
        run = paragraph.add_run(row['Modules not running'])
        if row['Modules not running']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,0,0)
        else:
            run.font.color.rgb = docx.shared.RGBColor(0,128,0)
        run.add_break()
        run = paragraph.add_run('I found the following timetable clashes:')
        run.add_break()
        run = paragraph.add_run(row['Timetable clashes'])
        if row['Timetable clashes']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,0,0)
        else:
            run.font.color.rgb = docx.shared.RGBColor(0,128,0)
        run.add_break()
        run = paragraph.add_run('I have the following comments for the adviser')
        run.add_break()
        run = paragraph.add_run(row['Adviser recommendations'])
        if row['Adviser recommendations']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,102,0)
        run.add_break()
        run.add_break()
        
        word_document.save(word_file_name)
   
def process_folder(folder_name):
    """Finds all student formfiles (all excel files) in a folder and performs advising checks on them
    
    Parameters:
    -----------
    
    folder_name : string
        The path to the folder with the files
        
    Returns:
    --------
    
    summary_data_frame : pandas data frame
        Data frame with one column per student. Contains the same columns as the data frame returned
        by process_form_file()
    """
    
    folder_entries = os.listdir(folder_name)
    
    form_files = []
    for entry in folder_entries:
        if (os.path.isfile(os.path.join(folder_name,entry)) and 
            (entry.endswith('.xlsx') or entry.endswith('.xltx')) and
             not entry.startswith('~$')): # this last bit is for when excel has the file open
            form_files.append(entry)
            
    if len(form_files) == 0:
        raise(ValueError('there are no forms in the folder you have given me'))
    else:
        list_of_data_frames = []
        for filename in form_files:
            this_data_frame = process_form_file(os.path.join(folder_name, filename))
            list_of_data_frames.append(this_data_frame)
            separation_string = '-'*60
            print(' ')
            print(separation_string)
            print(' ')
        summary_data_frame = pd.concat(list_of_data_frames, ignore_index=True)

        summary_data_frame = summary_data_frame.sort_values(by='Student ID')

        return summary_data_frame
 
from .programme_requirements import *
from .prerequisites import *
from .timetabling import *
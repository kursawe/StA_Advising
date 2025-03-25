import os
import openpyxl
import pandas as pd
import termcolor
from .student import *
import docx
import numbers
from datetime import date


module_catalogue_location = os.path.join(os.path.dirname(__file__),'Module_catalogue.xlsx') 
module_catalogue = pd.read_excel(module_catalogue_location)

def process_form_file_or_student_id(argument, programme_name = None):
    """preforms all advising checks on the 
    submitted form.
    
    Parameters:
    -----------
    
    argument : string or int
        if it's a string: path to the file that is being investigated,
        i.e. a filled-in module choice form
        if it's an integer: a valid student ID
        
    programme_name : string
        if this is not none than the programme requirements for this programme will be checked.
    """ 
    if isinstance(argument, str):
        student_or_warning = parse_excel_form(argument)
        _,filename_for_output = os.path.split(argument)
    elif isinstance(argument, numbers.Integral):
        student_or_warning = collect_student_data(argument)
        filename_for_output = 'student id ' + str(argument)
    else:
        raise(ValueError('Could not read argument of process_form_file_or_student, it is not an int or a string'))

    if isinstance(student_or_warning, str):
        if student_or_warning == 'No student ID':
            warning_message = 'Could not process ' + filename_for_output + '. The file does not contain a valid student ID.'
        elif student_or_warning.startswith('contains invalid student ID'):
            warning_message = 'Could not process ' + filename_for_output + '. The file ' + student_or_warning
        elif student_or_warning.startswith('Do not recognise student programme for parsing:'):
            warning_message = 'Could not process ' + filename_for_output + '. ' + student_or_warning
        colour_code_print_statement(warning_message)
        summary_data =[999999999, 
                    'Unknown',
                    'Unknown',
                    0,
                    warning_message, ' ', ' ', ' ', ' ']

        summary_data_frame = generate_summary_data_frame_from_entries(summary_data)

        return summary_data_frame
    
    student = student_or_warning
    if programme_name is not None:
        student.programme_name = programme_name
    print('Processing file or student')
    print(str(argument))
    print(' ')
    print('Student ID: ' + str(student.student_id))
    print('Name: ' + student.full_name)
    print('Programme: ' + student.programme_name)
    print(' ')
    
    missed_programme_requirements, programme_adviser_recommendations = find_missing_programme_requirements(student)
    
    print('The student is missing the following programme requirements:')
    colour_code_print_statement(missed_programme_requirements)

    missed_prerequisites, prerequisite_adviser_recommendations = find_missing_prerequisites(student)
    
    print('The student is missing the following prerequisites:')
    colour_code_print_statement(missed_prerequisites)

    not_running_modules, scheduling_adviser_recommendations = find_not_running_modules(student)

    print('The student selected the following modules when they are not running:')
    colour_code_print_statement(not_running_modules)

    timetable_clashes, timetable_adviser_recommendations = find_timetable_clashes(student)
    print('I found the following timetable clashes:')
    colour_code_print_statement(timetable_clashes)

    adviser_recommendations = merge_list_to_long_string([programme_adviser_recommendations, prerequisite_adviser_recommendations, 
                                                         scheduling_adviser_recommendations, timetable_adviser_recommendations])

    print('I have the following comments to the adviser:')
    colour_code_print_statement(adviser_recommendations, is_advice = True)

    summary_data = [student.student_id, 
                    student.full_name,
                    student.programme_name,
                    student.current_honours_year,
                    missed_programme_requirements, missed_prerequisites, not_running_modules, timetable_clashes, adviser_recommendations]

    summary_data_frame = generate_summary_data_frame_from_entries(summary_data)
    
    return summary_data_frame

def parse_excel_form(filename):
    """returns an instance of a 'student' class
    that has all the excel data as named attributes

    Parameters:
    -----------
    
    filename : string
        path to the file that is being investigated,
        i.e. a filled-in module choice form
        
    Returns:
    --------

    student : instance of Student class or string
        an object with student attributes. If it's a string then it's a warning
        message about what went wrong
    """
    # open the form file and read in the student ID
    this_workbook = openpyxl.load_workbook(filename=filename)
    sheet = this_workbook.active
    student_id = sheet["D5"].value
    if not isinstance(student_id, int):
        return 'No student ID'
        
    this_student = collect_student_data(student_id, include_credits=False)
    if isinstance(this_student, str):
        return this_student
   
    max_selected_honours_year_string = this_student.honours_module_choices['Honours year'].max()
    if pd.isna(max_selected_honours_year_string):
        first_year_to_read = this_student.current_honours_year
    else: 
        max_selected_honours_year = int(max_selected_honours_year_string[5])
        first_year_to_read = max_selected_honours_year + 1
    
    # read in modules for all honours years that have not happened yet

    module_table = []
    for remaining_honours_year in range(first_year_to_read,this_student.expected_honours_years + 1):
        year_key = 'Year ' + str(remaining_honours_year)
        calendar_year = this_student.current_calendar_year + remaining_honours_year - this_student.current_honours_year
        calendar_year_string = str(calendar_year) + '/' + str(calendar_year + 1)
        for semester_number in [1,2]:
            semester_modules = get_modules_under_header(sheet, year_key + ' of Honours: Semester ' + str(semester_number)) 
            for module in semester_modules:
                module_table.append([year_key, calendar_year_string, 'S' + str(semester_number), module,])

   # Turn this all into a nice pandas data frame
    honours_module_choices = pd.DataFrame(module_table, columns = ['Honours year', 'Academic year', 'Semester', 'Module code'])
    
    this_student.update_honours_module_choices(honours_module_choices)

    # return the student
    return this_student

def collect_student_data(student_id, include_credits = True):
    """Collects all available data for the student with the given ID
    
    Parameters :
    -----------
    
    student_id : int
        the student id
        
    Returns :
    ---------
    
    student : instance of Student class or string
        an object with student attributes. If it's a string then it's a warning
        message about what went wrong
    """
    # process data base here
    
    # Now that we have the student ID we can look up the student in the database:
    data_bases = get_all_mms_data_bases()

    # get a table with only the entries for this student
    student_not_yet_found = True

    data_base_index = 0
    while student_not_yet_found and data_base_index < len(data_bases):
        this_data_base = data_bases[data_base_index]
        if student_id in this_data_base['Student ID'].to_numpy():
            student_not_yet_found = False
            student_data_base = this_data_base[this_data_base['Student ID'] == student_id]
        data_base_index += 1
        
    if student_not_yet_found:
        return 'contains invalid student ID ' + str(student_id)

    # infer the year of study from the earliest module taken
    past_student_data_base = student_data_base.loc[(student_data_base['Assessment result'].notna()) &
                                                   (student_data_base['Assessment result'] != 'V') ] 
                                                #    (student_data_base['Semester'] == 'S2')]

    data_of_module_years = past_student_data_base['Year'].str.slice(0,4).astype('int')
    
    today = date.today()
    if today.month < 3:
        current_calendar_year = today.year - 1
    else:
        current_calendar_year = today.year

    if not pd.isna(data_of_module_years.min()):
        earliest_year = data_of_module_years.min()
        # current_calendar_year = data_of_module_years.max()+1
    else:
        earliest_year = current_calendar_year
        # earliest_year = 2024
        # current_calendar_year = 2024
    
    leave_of_absence_years = 0
    for year in range(earliest_year, current_calendar_year):
        if year not in data_of_module_years.to_list():
            leave_of_absence_years +=1
   
    # we add the +1 as the year of steady is 1-indexed instead of 0-indexed
    # i.e. students starting this year will be in year 1
    year_of_study = current_calendar_year - earliest_year + 1
    year_of_study = year_of_study - leave_of_absence_years
    
    # identify all modules that the student has passed
    data_base_of_passed_modules = student_data_base[(student_data_base['Assessment result']=='P') | 
                                                    (student_data_base['Reassessment result']=='P') |
                                                    ((student_data_base['Assessment result']=='Z') & (pd.isnull(student_data_base['Reassessment result']))) |
                                                    ((student_data_base['Assessment result']=='D') & (pd.isnull(student_data_base['Reassessment result']))) |
                                                    ((student_data_base['Assessment result']=='F') & (pd.isnull(student_data_base['Reassessment result'])) 
                                                                                                   & (student_data_base['Assessment grade'] > 3.5) ) | 
                                                    ( ( (student_data_base['Assessment result']=='S')|(student_data_base['Assessment result']=='SP') ) & 
                                                      ( (~pd.isnull(student_data_base['Assessment grade']) & (student_data_base['Assessment grade'] > 7.0) ) |
                                                          (pd.isnull(student_data_base['Reassessment result']))) )].copy()
    data_base_of_passed_modules.drop_duplicates(subset='Module code', keep='last', inplace=True)
    passed_modules = data_base_of_passed_modules['Module code'].to_list()
    passed_modules = passed_modules
    
    # remember all deferred and z coded modules
    data_base_of_z_coded_modules = student_data_base[(student_data_base['Assessment result']=='Z') & (pd.isnull(student_data_base['Reassessment result']))]
    z_coded_modules = data_base_of_z_coded_modules['Module code'].to_list()
    
    data_base_of_deferred_modules = student_data_base[(student_data_base['Assessment result']=='D') & (pd.isnull(student_data_base['Reassessment result']))]
    deferred_modules = data_base_of_deferred_modules['Module code'].to_list()
    
    data_base_of_modules_for_reassessment = data_base_of_passed_modules[(data_base_of_passed_modules['Assessment result']=='F') & (pd.isnull(data_base_of_passed_modules['Reassessment result'])) 
                                                                                                   & (data_base_of_passed_modules['Assessment grade'] > 3.5) ]
    modules_awaiting_reassessment = data_base_of_modules_for_reassessment['Module code'].to_list()

    data_base_of_s_coded_modules = student_data_base[( (student_data_base['Assessment result']=='S')|(student_data_base['Assessment result']=='SP') ) & 
                                                     ( (~pd.isnull(student_data_base['Assessment grade']) & (student_data_base['Assessment grade'] < 7.0) ) &
                                                          (pd.isnull(student_data_base['Reassessment result']))) ]

    s_coded_modules = data_base_of_s_coded_modules['Module code'].to_list()

    #identify the programme of the student
    programme_entries = student_data_base['Programme name'].unique()    
    assert(len(programme_entries) == 1)
    programme_name = programme_entries[0]
    
    given_name_entries = student_data_base['Given names'].unique()
    assert(len(given_name_entries) == 1)
    given_name = given_name_entries[0]
    
    family_name_entries = student_data_base['Family name'].unique()
    assert(len(family_name_entries) == 1)
    family_name = family_name_entries[0]
    
    full_name = given_name + ' ' + family_name

    email_entries = student_data_base['Email'].unique()
    assert(len(email_entries) == 1)
    email = email_entries[0]

    # Figure out what year they are in and how many they have left
    if 'Bachelor of Science' in programme_name:
        no_of_programme_years = 4
        expected_honours_years = 2
    elif 'Master in Mathematics' in programme_name:
        no_of_programme_years = 5
        expected_honours_years = 3
    elif 'Master of Arts (Honours)' in programme_name: 
        no_of_programme_years = 4
        expected_honours_years = 2
    elif 'Master in Chemistry (Honours) Chemistry with Mathematics' == programme_name:
        no_of_programme_years = 5
        expected_honours_years = 3
    elif 'Master in Physics (Honours) Mathematics and Theoretical Physics' == programme_name:
        no_of_programme_years = 5
        expected_honours_years = 3
    else:
        warning_message = 'Do not recognise student programme for parsing: ' + programme_name
        return warning_message
    
    if 'EXA120' in student_data_base['Module code'].values:
        no_of_programme_years -=1
        
    no_subhonours_years = no_of_programme_years - expected_honours_years
    current_honours_year = year_of_study - no_subhonours_years
    
    # make a separate data base of passed honours modules
    passed_honours_modules = list()
    for previous_honours_year in range(1,current_honours_year+1):
        year_difference = current_honours_year - previous_honours_year  
        year_number = current_calendar_year-year_difference
        calendar_year_string = str(year_number) + '/' + str(year_number + 1)
        data_base_of_passed_modules_this_year = data_base_of_passed_modules[data_base_of_passed_modules['Year']==calendar_year_string]
        passed_modules_this_year = data_base_of_passed_modules_this_year['Module code'].to_list()
        passed_honours_modules += passed_modules_this_year
        
    # updating first honours year if we know when the student has taken the first MT3* module 
    student_data_base_with_valid_module_codes = student_data_base[student_data_base['Module code'].notna()]
    data_base_of_honours_modules = student_data_base_with_valid_module_codes[student_data_base_with_valid_module_codes['Module code'].str.contains('MT3')]
    if not data_base_of_honours_modules.empty:
        data_of_honours_module_years = data_base_of_honours_modules['Year'].str.slice(0,4).astype('int')
        first_honours_year = data_of_honours_module_years.min()
        current_honours_year = current_calendar_year - first_honours_year + 1
        leave_of_absence_years_honours = 0
        for year in range(first_honours_year, current_calendar_year):
            if year not in data_of_module_years.to_list():
                leave_of_absence_years_honours +=1
        current_honours_year -= leave_of_absence_years_honours
        year_of_study = current_honours_year + no_subhonours_years
        # use this information to get a more accurate guess of what honours modules they have taken
        passed_honours_modules = list()
        for previous_year in range(first_honours_year,current_calendar_year+1):
            calendar_year_string = str(previous_year) + '/' + str(previous_year + 1)
            data_base_of_passed_modules_this_year = data_base_of_passed_modules[data_base_of_passed_modules['Year']==calendar_year_string]
            passed_modules_this_hear = data_base_of_passed_modules_this_year['Module code'].to_list()
            passed_honours_modules += passed_modules_this_hear

    passed_module_table = reduce_official_data_base(data_base_of_passed_modules, current_honours_year, current_calendar_year) 
    
    data_base_of_planned_modules = student_data_base[(pd.isna(student_data_base['Assessment result']))]
    honours_module_choices = reduce_official_data_base(data_base_of_planned_modules, current_honours_year, current_calendar_year)
    
    # finish processing forms

    # module_table= []
    # honours_module_choices = pd.DataFrame(module_table, columns = ['Honours year', 'Academic year', 'Semester', 'Module code'])
    this_student = Student(student_id, 
                           full_name,
                           email,
                           programme_name, 
                           year_of_study,
                           expected_honours_years,
                           current_honours_year,
                           current_calendar_year,
                           passed_modules,
                           z_coded_modules,
                           deferred_modules,
                           s_coded_modules,
                           modules_awaiting_reassessment,
                           passed_module_table,
                           passed_honours_modules,
                           honours_module_choices)
    
    return this_student
 
def get_all_mms_data_bases():
    '''looks into the student_data folder and loads all data bases in there that it can find into memory.
    
    Returns :
    ---------
    
    data_bases: list of pandas data frames
        each entry of the list is one pandas data frame from a .csv file found in the student_data folder.
    '''
    # Now that we have the student ID we can look up the student in the database:
    current_file_directory = os.path.dirname(os.path.abspath(__file__))

    # Find the data directory; this should be at ../../student_data (if running advising_tool.py)
    # or at the CWD[/student_data] if using as a library
    path_if_tool = os.path.join(current_file_directory, '../..', 'student_data')
    path_if_cwd = os.path.join(os.getcwd(), "student_data")

    if os.path.exists(path_if_tool):
        data_directory = path_if_tool
    elif os.path.exists(os.path.join(os.getcwd(), "student_data")):
        data_directory = path_if_cwd
    else:
        data_directory = os.getcwd()
    potential_data_files = os.listdir(data_directory)
    data_files = []
    for candidate_filename in potential_data_files:
       this_filename, file_extension = os.path.splitext(candidate_filename)
       if file_extension == '.csv':
           data_files.append(candidate_filename)
    if len(data_files) == 0:
        raise FileNotFoundError("missing student data .csv file")
    
    # turn them all into data base files
    data_bases = []
    for data_file_name in data_files:
        data_path = os.path.join(data_directory, data_file_name)
        this_data_frame = pd.read_csv(data_path)
        this_data_frame = this_data_frame.map(strip_excel_formatting)
        data_bases.append(this_data_frame.astype({"Student ID": "int64",
                                                  "Credits": "float64"}))

    return data_bases

# Some data files come with data in the form `="..."`; strip this if it exists
def strip_excel_formatting(cell_data):
    if isinstance(cell_data, str):
        if cell_data[:2] == '="' and cell_data[-1] == '"':
            return cell_data[2:-1]
    return cell_data


def reduce_official_data_base(data_frame, current_honours_year, current_calendar_year, include_credits = True):
    '''take a table from the official data base and reduce it to a smaller pandas data frame that only has entries that
    we care about.
    
    Parameters :
    ------------
    
    data_frame : pandas dataframe
        this must have been constructed from one of the official MMS data base files
        
    current_honours_year : int
        the current honours year of the student being processed
        
    current_honours_year : int
        the current calendar year

    Returns :
    ---------
    
    reduced_data_frame : pandas data frame
        will only contains essential module information
    '''
    module_table = []
    # past_data_frame = data_frame.loc[(data_frame['Assessment result'].notna()) &
                                    #  (data_frame['Semester'] == 'S2')]

    # if not past_data_frame.empty:
        # data_of_module_years = past_data_frame['Year'].str.slice(0,4).astype('int')
        # most_recent_year = data_of_module_years.max() +1
    # else:
        #in this case, the data frame we are given only has selected, not past modules
        #hence, the minimal selected year must be the current actual year
        # data_of_module_years = data_frame['Year'].str.slice(0,4).astype('int')
        # most_recent_year = data_of_module_years.min()
    most_recent_year = current_calendar_year

    # data_of_module_years = past_student_data_base['Year'].str.slice(0,4).astype('int')
    
    for _,row in data_frame.iterrows():
        module_code = row['Module code']
        academic_year = row['Year']
        semester = row['Semester']
        credits = row['Credits available']
        this_year = int(academic_year[:4])
        this_honours_year =  this_year - most_recent_year + current_honours_year
        this_honours_year_string = 'Year ' + str(this_honours_year)
        if include_credits:
            module_table.append([this_honours_year_string, 
                                    academic_year, semester, module_code,credits,])
        else:
            module_table.append([this_honours_year_string, 
                                    academic_year, semester, module_code,])
            

    # Turn this all into a nice pandas data frame
    if include_credits:
        reduced_data_frame = pd.DataFrame(module_table, 
                                           columns = ['Honours year', 'Academic year', 'Semester', 'Module code', 'Credits'])
    else:
        reduced_data_frame = pd.DataFrame(module_table, 
                                           columns = ['Honours year', 'Academic year', 'Semester', 'Module code'])

    return reduced_data_frame

def generate_summary_data_frame_from_entries(data_list):
    """Generates the summary data frame from the entries for one student.
    
    Parameters:
    -----------

    data_list : list
        list entries vary in type, see the expected order of entries in the code below.
        
    Returns :
    ---------
    
    summary_data_frame : pandas data frame
        the pandas data frame containing the list data and the correct headers.
    """
    summary_data_frame = pd.DataFrame([data_list], columns = ['Student ID', 
                                                                 'Name',
                                                                 'Programme',
                                                                 'Hon. year',
                                                                 'Unmet programme requirements', 'Missing prerequisites', 'Modules not running', 'Timetable clashes', 'Adviser recommendations'])
    
    return summary_data_frame
   
def colour_code_print_statement(print_statement, is_advice = False):
    """Prints the given string in red if the string is not equal to 'None'.
    Otherwise it will print in green.
    
    Parameters:
    -----------
    
    print_statement : string
        the statement to be printed
        
    is_advice : bool
        if False, strings that are not 'None' will be printed in light magenta instead of red
    """
    if print_statement != 'None':
        if is_advice:
            print(termcolor.colored(print_statement,'blue'))
        else:
            print(termcolor.colored(print_statement,'red'))
    else:
        print(termcolor.colored('None','green'))
        
    #add an empty line by default
    print(' ')

def merge_list_to_long_string(a_list):
    """takes a list of strings and returns a string that separates the entries with a comma and a space.
        Returns the string 'None' if the list is empty. If the first entry of the list is 'None' then that entry will be ignored.
    
    Parameters:
    -----------
    a_list : list
        the list we want to parse, for example a list of unmet programme requirements
    
    Returns:
    --------
    a_string : string
        contains all entries in a_list separated by a comma and a space
        is 'None' if a_list is empty
    """
    a_string = ''
    first_item_found = False
    for item in a_list:
        if item != 'None':
            if first_item_found:
                a_string += '\n' + item
            else:
                a_string = item
                first_item_found = True
    
    if a_string == '':
        a_string = 'None'
    
    return a_string
 
def get_modules_under_header(sheet, header):
    """get all the modules in the student module choice form under a given heading
    
    Parameters:
    -----------
    
    sheet : openpyxl sheet object
        generated from the workbook excel file
        
    header : string
        the header that we are investigating

    Returns:
    --------

    modules : list of strings
        module codes under the given header
    """
    modules = []
    for row in sheet.iter_rows():
        if isinstance(row[1].value, str):
            if header in row[1].value:
                row_number = row[1].row

    # the row number where modules are entered is 2 down
    row_number += 2
    next_cell_name_with_module = 'B' + str(row_number)
    module_code_is_not_empty = sheet[next_cell_name_with_module].value is not None

    row_index = 0
    while row_index < 6:
        this_cell_is_empty = sheet[next_cell_name_with_module].value is None
        if not this_cell_is_empty:
            module_code_entry = sheet[next_cell_name_with_module].value
            if isinstance(module_code_entry, int):
                module_code = str(module_code_entry)
                module_code = 'MT' + module_code
            else:
                module_code = module_code_entry
            module_code = module_code.strip()
            module_code = module_code.replace('Mt','MT')
            modules.append(module_code)
        row_number+=1 
        next_cell_name_with_module = 'B' + str(row_number)
        # module_code_is_not_empty = sheet[next_cell_name_with_module].value is not None
        row_index += 1

    return modules

def colour_code_passes(column):
    """Helper function for saving the summary data frame, will colour code the column entries in a 
    pandas array in green or red depending on whether they contain the entry 'None' or not
    
    Parameters:
    -----------

    column: pandas data array column
       the column we are colour coding
    
    Returns:
    --------

    color_list : list of strings
       a list of color setting for each of the cells in the array
    """
    passed_colour = 'background-color: palegreen;'
    failed_colour = 'background-color: lightcoral;'
    return [passed_colour if value=='None' else failed_colour for value in column]

def colour_recommendations(column):
    """Helper function for saving the summary data frame, will colour code the column entries in a 
    pandas array in no color or orange depending on whether they contain the entry 'None' or not
    
    Parameters:
    -----------

    column: pandas data array column
       the column we are colour coding
    
    Returns:
    --------

    color_list : list of strings
       a list of color setting for each of the cells in the array
    """
    recommendation_colour = 'background-color: orange;'
    default_colour = ''
    return [default_colour if value=='None' else recommendation_colour for value in column]

def save_summary_data_frame(data_frame, filename):
    """Save the summary data frame two an excel file, inlcuding a colour code formatting 
    for the column.
    
    Parameters :
    ------------

    data_frame : Pandas data frame,
        typically generated from the 'check_form_file' function
    
    filename : string
        where we want to save
    """

    if not filename.endswith('.xlsx'):
        if filename.endswith('.docx'):
            filename = filename[:-5] + '.xlsx'
        else:
            filename += '.xlsx'
        
    # Let's start by writing into word 
    filename_base = filename[:-5]
    word_file_name = filename_base + '.docx'
    
    word_document = docx.Document()
    
    for _,row in data_frame.iterrows():
        paragraph = word_document.add_paragraph()
        run = paragraph.add_run('Student ID: ' + str(row['Student ID']))
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run('Name: ' + row['Name'])
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run('Programme: ' + row['Programme'])
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run('The student is missing the following programme requirements:')
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run(row['Unmet programme requirements'])
        if row['Unmet programme requirements']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,0,0)
        else:
            run.font.color.rgb = docx.shared.RGBColor(0,128,0)
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run('The student is missing the following prerequisites:')
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run(row['Missing prerequisites'])
        if row['Missing prerequisites']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,0,0)
        else:
            run.font.color.rgb = docx.shared.RGBColor(0,128,0)
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run('The student selected the following modules when they are not running:')
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run(row['Modules not running'])
        if row['Modules not running']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,0,0)
        else:
            run.font.color.rgb = docx.shared.RGBColor(0,128,0)
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run('I found the following timetable clashes:')
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run(row['Timetable clashes'])
        if row['Timetable clashes']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,0,0)
        else:
            run.font.color.rgb = docx.shared.RGBColor(0,128,0)
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run('I have the following comments for the adviser')
        run.font.name='Arial'
        run.add_break()
        run = paragraph.add_run(row['Adviser recommendations'])
        if row['Adviser recommendations']!= 'None':
            run.font.color.rgb = docx.shared.RGBColor(255,102,0)
        run.font.name='Arial'
        run.add_break()
        run.add_break()
        
        word_document.save(word_file_name)
 
    text_columns = ['Unmet programme requirements', 'Missing prerequisites', 'Modules not running', 'Timetable clashes', 'Adviser recommendations']
    data_frame[text_columns] = data_frame[text_columns].replace('\n', '; ', regex=True)
    styled_data_frame = (data_frame.style.apply(colour_code_passes, subset = ['Unmet programme requirements', 'Missing prerequisites', 'Modules not running',
                                                                                       'Timetable clashes'], axis = 0).
                          apply(colour_recommendations, subset = ['Adviser recommendations'], axis = 0))


    ## write to excel

    writer = pd.ExcelWriter(filename) 
    # Manually adjust the width of each column
    styled_data_frame.to_excel(writer)
    worksheet = writer.sheets['Sheet1']
    font = openpyxl.styles.Font(size=14)
    worksheet.set_column(0,0,width=5)
    worksheet.set_column(1,1,width=12)
    worksheet.set_column(2,2,width=22)
    worksheet.set_column(3,3,width=35)
    worksheet.set_column(4,4,width=8)
    worksheet.set_column(5,5,width=40)
    worksheet.set_column(6,6,width=40)
    worksheet.set_column(7,7,width=40)
    worksheet.set_column(8,8,width=40)
    worksheet.set_column(9,9,width=40)
    
    # The only way I found to change the font size is to save the excel file, reload it, and then edit the fontsize, and save it again
    writer.close()
    
    # reload
    reloaded_workbook = openpyxl.load_workbook(filename)

    # Access the active sheet
    worksheet = reloaded_workbook.active

    # Set the font size for all cells in the worksheet
    font_size = 14
    font = openpyxl.styles.Font(size=font_size)
    for row in worksheet.iter_rows():
        for cell in row:
            cell.font = font

    # Make the first row bold
    bold_font = openpyxl.styles.Font(size=font_size, bold=True)
    for cell in worksheet[1]:
        cell.font = bold_font
    # Save the modified workbook
    reloaded_workbook.save(filename)
    
  
def process_folder(folder_name, programme_name = None):
    """Finds all student formfiles (all excel files) in a folder and performs advising checks on them
    
    Parameters:
    -----------
    
    folder_name : string
        The path to the folder with the files
        
    Returns:
    --------
    
    summary_data_frame : pandas data frame
        Data frame with one column per student. Contains the same columns as the data frame returned
        by process_form_file_or_student_id()


    programme_name : string
        If this is not None than the programme requirements for this programme will be checked.
        Note that every student in the folder will be checked against these requirements.
    """
    
    folder_entries = os.listdir(folder_name)
    
    form_files = []
    for entry in folder_entries:
        if (os.path.isfile(os.path.join(folder_name,entry)) and 
            (entry.endswith('.xlsx') or entry.endswith('.xltx')) and
             not entry.startswith('~$')): # this last bit is for when excel has the file open
            form_files.append(entry)
            
    if len(form_files) == 0:
        raise(ValueError('there are no forms in the folder you have given me'))
    else:
        list_of_data_frames = []
        for filename in form_files:
            this_data_frame = process_form_file_or_student_id(os.path.join(folder_name, filename), programme_name)
            list_of_data_frames.append(this_data_frame)
            separation_string = '-'*60
            print(' ')
            print(separation_string)
            print(' ')
        summary_data_frame = pd.concat(list_of_data_frames, ignore_index=True)

        summary_data_frame = summary_data_frame.sort_values(by='Student ID')

        return summary_data_frame
    
def check_final_year_students():
    '''
    Go through the data base, identify final year students, and check their
    module choices against the programme requirements

    Returns:
    --------
    
    summary_data_frame : pandas data frame
        Data frame with one column per student. Contains the same columns as the data frame returned
        by process_form_file()
    '''
    # turn them all into data base files
    data_bases = get_all_mms_data_bases()
    processed_students = []
    list_of_summary_data_frames = []
    for this_data_base in data_bases:
        these_student_ids = this_data_base['Student ID'].unique()
        for student_id in these_student_ids:
            if student_id not in processed_students:
                processed_students.append(student_id)
                student_or_warning = collect_student_data(student_id)
                if isinstance(student_or_warning, str):
                    colour_code_print_statement(student_or_warning)
                else:
                    student= student_or_warning
                    if student.current_honours_year >= student.expected_honours_years:
                        this_summary_data_frame = process_form_file_or_student_id(student_id)
                        list_of_summary_data_frames.append(this_summary_data_frame)
                        separation_string = '-'*60
                        print(' ')
                        print(separation_string)
                        print(' ')

    summary_data_frame = pd.concat(list_of_summary_data_frames, ignore_index=True)

    summary_data_frame_sorted = summary_data_frame.sort_values(by='Student ID')
    
    summary_data_frame = summary_data_frame_sorted[~summary_data_frame_sorted['Unmet programme requirements'].str.contains("No programme requirements available", na=False)]

    return summary_data_frame

from .programme_requirements import *
from .prerequisites import *
from .timetabling import *